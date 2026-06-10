#!/usr/bin/env python3
"""Generate WasiMahin_Resume.docx + .pdf from one content structure.

ATS rules baked in: single column, standard headings, no tables/images,
real bullets, Letter size, content-driven length (1-2 pages, never cram).
Run from the website repo root:  python3 .claude/resume_gen.py
"""

# ─────────────────────────── CONTENT ───────────────────────────

NAME = "Wasi Mahin"
CONTACT = ("Long Beach, CA  |  wasimahin@gmail.com  |  linkedin.com/in/wasi-mahin"
           "  |  github.com/wasimahin  |  wasimahin.com")

SUMMARY = (
    "Data analyst and finance leader pursuing a B.S. in Management Information Systems "
    "& Accountancy at CSULB (GPA 3.87). I solve operational problems with data, systems, "
    "and technology: engineered two production Python automation tools running daily inside "
    "Microsoft Dynamics 365 Business Central (~96% and 95% processing-time reductions), "
    "built Power BI dashboards and ETL pipelines for a Microsoft-affiliated project, and "
    "support financial operations for 500+ campus organizations using SQL, Excel, and "
    "Microsoft Copilot."
)

EXPERIENCE = [
    {
        "role": "Lead Business Office Representative",
        "org": "Associated Students, Inc., CSULB", "loc": "Long Beach, CA",
        "dates": "Feb 2026 – Present",
        "bullets": [
            "Engineered two production Python desktop automation tools, both live in daily "
            "operations with zero API integration: Snehin Check Request Helper (pdfplumber, "
            "PyMuPDF, pyautogui) parses vendor invoice PDFs and auto-enters data into Microsoft "
            "Dynamics 365 Business Central, cutting per-invoice processing from 3–5 minutes to "
            "under 10 seconds (~96%); Business Central Journal Entry Bot (pdfplumber, openpyxl, "
            "CustomTkinter) extracts journal entries from DocuSign PDFs and CashNet Excel files "
            "with debit/credit structuring and object-code normalization (95% reduction).",
            "Oversee financial documentation for 500+ registered campus organizations — accounts "
            "payable, purchase orders, bank reconciliations, document auditing, and compliance "
            "verification.",
            "Automated data extraction, validation, and reporting pipelines using Microsoft Copilot "
            "and prompt engineering, reducing processing time 70–80% across ASI and USU operations.",
            "Deliver SQL-driven budget analysis supporting ASI and USU budget creation; perform "
            "financial statement analysis and auditing across all 500+ accounts, maintaining 100% "
            "data accuracy across reporting cycles.",
            "Redesigned end-to-end office workflows, increasing operational efficiency 50% — "
            "standardized protocols formally adopted team-wide; mentor team members on financial "
            "systems and compliance.",
        ],
    },
    {
        "role": "Vice President of Finance",
        "org": "Association for Information Systems, CSULB", "loc": "Long Beach, CA",
        "dates": "May 2026 – Present",
        "bullets": [
            "Own the full financial lifecycle of a 100+ member organization: $5,000+ operating "
            "budget and external grant funding across weekly workshops, seminars, and semester events.",
            "Analyze financial data to optimize fund distribution; drive improvements to financial "
            "tracking, reporting, and allocation processes for efficiency, transparency, and scalability.",
            "Collaborate with the executive board to align financial strategy with membership growth "
            "and program development goals.",
        ],
    },
    {
        "role": "Business Analyst Intern (Project Microsoft)",
        "org": "Association for Information Systems, CSULB", "loc": "Long Beach, CA",
        "dates": "Sep 2025 – Nov 2025",
        "bullets": [
            "Engineered interactive Power BI dashboards using DAX and custom data models, "
            "transforming raw audience and performance datasets into executive-ready KPI "
            "visualizations for Microsoft presentations.",
            "Built ETL pipelines using Microsoft Fabric and Copilot; conducted correlation analysis "
            "and market research that directly shaped project strategy and implementation decisions.",
            "Executed SWOT and cost-benefit feasibility analyses; co-authored and delivered business "
            "pitches to Microsoft representatives, securing stakeholder confidence.",
            "Coordinated cross-departmental data validation and cleansing, enforcing accuracy "
            "standards and reducing downstream reporting errors.",
        ],
    },
    {
        "role": "Business Office Representative",
        "org": "Associated Students, Inc., CSULB", "loc": "Long Beach, CA",
        "dates": "Oct 2024 – Feb 2026",
        "bullets": [
            "Reduced check processing time 80% through implementation of an upgraded financial "
            "system, improving throughput during high-volume transaction periods.",
            "Automated data extraction, transformation, and visualization pipelines using Microsoft "
            "Copilot, cutting processing time 70% and freeing capacity for analytical work.",
            "Co-developed the 2025–2026 ASI and USU operating budgets; managed and reconciled "
            "financial statements for 20+ student organizations.",
            "Collaborated cross-functionally to verify accounting entries, maintaining data accuracy "
            "and preventing downstream reporting errors.",
        ],
    },
    {
        "role": "Junior Accountant",
        "org": "Mizan Enterprise", "loc": "Dhaka, Bangladesh",
        "dates": "Feb 2023 – Jul 2023",
        "bullets": [
            "Managed and reconciled daily expenses in Excel, ensuring 100% ledger accuracy.",
            "Created analytical reports highlighting financial discrepancies, improving reporting "
            "processes 15%.",
        ],
    },
]

PROJECTS = [
    ("Snehin Check Request Helper",
     "Python RPA tool in daily production at ASI — github.com/wasimahin/snehin-check-request-helper"),
    ("Business Central Journal Entry Bot",
     "Python journal-entry automation in daily production — github.com/wasimahin/bc-journal-entry-bot"),
    ("“Is the American Dream Still Attainable?”",
     "Tableau economic analysis of income growth and regional inequality — 3rd place, "
     "AIS Annual Datathon 2026."),
]

EDUCATION = [
    ("California State University, Long Beach",
     "B.S., Management Information Systems & Accountancy — Expected May 2027",
     "GPA 3.87  |  President's Honor List, every semester since Fall 2023"),
]

SKILLS = [
    ("Languages & Querying", "Python, SQL, DAX, Excel (Power Query, Pivot Tables, VLOOKUP), pdfplumber, openpyxl"),
    ("BI & Visualization", "Power BI, Tableau, Data Modeling, KPI Dashboards, Data Storytelling"),
    ("Data Engineering & AI", "ETL Pipelines, Microsoft Fabric, Microsoft Copilot, Prompt Engineering, Workflow Automation, Data Validation"),
    ("Finance", "Financial Analysis, Budgeting, Forecasting, Accounts Payable, Bank Reconciliation, Variance Analysis"),
    ("Platforms", "Microsoft Dynamics 365 Business Central, Microsoft 365, GitHub"),
]

CERTIFICATIONS = [
    "Career Essentials in Data Analysis — Microsoft & LinkedIn",
    "Data Manipulation in SQL — DataCamp",
    "Creating & Analyzing Dashboards in Tableau — DataCamp",
    "Analyzing Data in Tableau — DataCamp",
    "Introduction to Tableau — DataCamp",
]

# ─────────────────────────── DOCX ───────────────────────────

def build_docx(path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(0.7)
        sec.top_margin = sec.bottom_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(3)

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        # thin bottom border for visual separation (ATS-safe paragraph border)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "444444")
        pbdr.append(bottom); pPr.append(pbdr)

    # Header
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(NAME); r.bold = True; r.font.size = Pt(20)
    p = doc.add_paragraph(CONTACT); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)

    heading("Summary")
    doc.add_paragraph(SUMMARY)

    heading("Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(job["role"]); r.bold = True
        p.add_run(f"  —  {job['org']}  |  {job['loc']}  |  ").italic = False
        r = p.add_run(job["dates"]); r.italic = True
        for b in job["bullets"]:
            bp = doc.add_paragraph(b, style="List Bullet")
            bp.paragraph_format.space_after = Pt(2)

    heading("Projects")
    for title, desc in PROJECTS:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(title + " — "); r.bold = True
        p.add_run(desc)

    heading("Education")
    for school, degree, detail in EDUCATION:
        p = doc.add_paragraph()
        r = p.add_run(school); r.bold = True
        doc.add_paragraph(degree)
        doc.add_paragraph(detail)

    heading("Skills")
    for group, items in SKILLS:
        p = doc.add_paragraph()
        r = p.add_run(group + ": "); r.bold = True
        p.add_run(items)

    heading("Certifications")
    for cert in CERTIFICATIONS:
        doc.add_paragraph(cert, style="List Bullet")

    doc.save(path)
    print("wrote", path)

# ─────────────────────────── PDF ───────────────────────────

def build_pdf(path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable, ListFlowable, ListItem)

    black = HexColor("#111111"); gray = HexColor("#444444")
    s_name = ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=19,
                            leading=23, alignment=TA_CENTER, textColor=black)
    s_contact = ParagraphStyle("contact", fontName="Helvetica", fontSize=9.5,
                               leading=13, alignment=TA_CENTER, textColor=gray)
    s_h = ParagraphStyle("h", fontName="Helvetica-Bold", fontSize=11.5,
                         leading=14, spaceBefore=10, spaceAfter=2, textColor=black)
    s_body = ParagraphStyle("body", fontName="Helvetica", fontSize=9.8,
                            leading=13.2, textColor=black, spaceAfter=3)
    s_role = ParagraphStyle("role", parent=s_body, spaceBefore=6, spaceAfter=2)
    s_bullet = ParagraphStyle("bullet", parent=s_body, spaceAfter=2)

    doc = SimpleDocTemplate(path, pagesize=letter,
                            leftMargin=.7*inch, rightMargin=.7*inch,
                            topMargin=.55*inch, bottomMargin=.55*inch,
                            title=f"{NAME} — Resume", author=NAME)
    story = [Paragraph(NAME, s_name), Paragraph(CONTACT, s_contact)]

    def heading(text):
        story.append(Paragraph(text.upper(), s_h))
        story.append(HRFlowable(width="100%", thickness=0.7, color=gray, spaceAfter=5))

    def bullets(items):
        story.append(ListFlowable(
            [ListItem(Paragraph(b, s_bullet), leftIndent=14) for b in items],
            bulletType="bullet", start="•", bulletFontSize=9, leftIndent=14))

    heading("Summary")
    story.append(Paragraph(SUMMARY, s_body))

    heading("Experience")
    for job in EXPERIENCE:
        story.append(Paragraph(
            f"<b>{job['role']}</b> — {job['org']}  |  {job['loc']}  |  <i>{job['dates']}</i>",
            s_role))
        bullets(job["bullets"])

    heading("Projects")
    bullets([f"<b>{t}</b> — {d}" for t, d in PROJECTS])

    heading("Education")
    for school, degree, detail in EDUCATION:
        story.append(Paragraph(f"<b>{school}</b>", s_role))
        story.append(Paragraph(degree, s_body))
        story.append(Paragraph(detail, s_body))

    heading("Skills")
    for group, items in SKILLS:
        story.append(Paragraph(f"<b>{group}:</b> {items}", s_body))

    heading("Certifications")
    bullets(CERTIFICATIONS)

    doc.build(story)
    print("wrote", path)

if __name__ == "__main__":
    build_docx("WasiMahin_Resume.docx")
    build_pdf("WasiMahin_Resume.pdf")
